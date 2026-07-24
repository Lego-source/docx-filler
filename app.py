import base64
import io
import re
import zipfile
import copy
from flask import Flask, request, Response

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.enum.text import WD_COLOR_INDEX

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024

APP_VERSION = "2026-07-24-leadership-v11-genitive-head"


def normalize_name(s):
    s = (s or '')
    s = s.replace('_', ' ').replace('\u00A0', ' ')
    s = re.sub(r'\s+', ' ', s).strip()
    return s.lower()


def normalize_apostrophes_(s):
    return (s or '').replace('\u2019', "'").replace('\u02bc', "'")


def normalize_for_anchor_(s):
    s = (s or '').lower()
    s = s.replace('\u00A0', ' ')
    for ch in ["'", '\u2019', '\u02bc', '\u0060', '\u00B4']:
        s = s.replace(ch, '')
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def build_matcher(data_map):
    lookup = {}
    names = []
    for key, val in data_map.items():
        norm = normalize_name(key)
        if norm:
            lookup[norm] = '' if val is None else str(val)
            names.append(key)
    if not names:
        return None, {}
    names.sort(key=lambda n: len(n), reverse=True)
    alts = []
    for n in names:
        piece = re.escape(n.strip())
        piece = re.sub(r'\\?\s+', r'[\\s_]+', piece)
        alts.append(piece)
    pattern = re.compile(
        '\u00AB[\\s_]*(' + '|'.join(alts) + ')[\\s_]*\u00BB',
        re.IGNORECASE | re.UNICODE
    )
    return pattern, lookup


def collect_paragraphs(container):
    paras = []
    try:
        paras.extend(container.paragraphs)
    except Exception:
        pass
    try:
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras.extend(collect_paragraphs(cell))
    except Exception:
        pass
    return paras


def collect_paragraphs_dedup(container):
    paras = []
    seen = set()

    def _walk(c):
        try:
            for p in c.paragraphs:
                pid = id(p._p)
                if pid not in seen:
                    seen.add(pid)
                    paras.append(p)
        except Exception:
            pass
        try:
            for table in c.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _walk(cell)
        except Exception:
            pass

    _walk(container)
    return paras


def collect_all_paragraphs_incl_headers(doc):
    result = list(collect_paragraphs_dedup(doc))
    for container in all_containers(doc):
        if container is doc:
            continue
        try:
            result.extend(collect_paragraphs_dedup(container))
        except Exception:
            pass
    return result


def walk_all_tables(container):
    tables = []
    try:
        for t in container.tables:
            tables.append(t)
            for row in t.rows:
                for cell in row.cells:
                    tables.extend(walk_all_tables(cell))
    except Exception:
        pass
    return tables


def cell_plain_text(cell):
    return ' '.join(''.join(r.text for r in p.runs) for p in cell.paragraphs)


def all_containers(doc):
    containers = [doc]
    for section in doc.sections:
        containers.append(section.header)
        containers.append(section.first_page_header)
        containers.append(section.even_page_header)
        containers.append(section.footer)
        containers.append(section.first_page_footer)
        containers.append(section.even_page_footer)
    return containers


def full_doc_text_lower(all_paragraphs):
    parts = [''.join(r.text for r in p.runs) for p in all_paragraphs]
    return normalize_for_anchor_(' '.join(parts))


def replace_in_paragraph(paragraph, pattern, lookup):
    runs = paragraph.runs
    if not runs:
        return
    if '\u00AB' not in ''.join(r.text for r in runs):
        return
    guard = 0
    while True:
        guard += 1
        if guard > 2000:
            break
        texts = [r.text for r in runs]
        full = ''.join(texts)
        m = pattern.search(full)
        if not m:
            break
        inner = normalize_name(m.group(1))
        value = lookup.get(inner, '')
        s, e = m.start(), m.end()
        offsets = []
        pos = 0
        for t in texts:
            offsets.append(pos)
            pos += len(t)
        first_i = None
        for i in range(len(texts)):
            if offsets[i] <= s < offsets[i] + len(texts[i]):
                first_i = i
                break
        last_i = None
        for j in range(len(texts)):
            if offsets[j] < e <= offsets[j] + len(texts[j]):
                last_i = j
        if first_i is None:
            break
        if last_i is None:
            last_i = len(texts) - 1
        prefix = texts[first_i][:s - offsets[first_i]]
        suffix = texts[last_i][e - offsets[last_i]:]
        if first_i == last_i:
            runs[first_i].text = prefix + value + suffix
        else:
            runs[first_i].text = prefix + value
            for k in range(first_i + 1, last_i):
                runs[k].text = ''
            runs[last_i].text = suffix


def apply_placeholders(doc, data_map):
    pattern, lookup = build_matcher(data_map)
    if pattern is None:
        return
    for p in collect_paragraphs(doc):
        replace_in_paragraph(p, pattern, lookup)
    for section in doc.sections:
        for hf in (section.header, section.first_page_header, section.even_page_header,
                   section.footer, section.first_page_footer, section.even_page_footer):
            try:
                for p in collect_paragraphs(hf):
                    replace_in_paragraph(p, pattern, lookup)
            except Exception:
                pass


def insert_paragraph_after(anchor_paragraph, text):
    new_p = copy.deepcopy(anchor_paragraph._p)
    anchor_paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor_paragraph._parent)
    runs = new_para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ''
    else:
        new_para.add_run(text)
    return new_para


TVO_A_MARK   = 'прошу покласти тимчасове виконання'
TVO_B_START  = 'не заперечую'
TVO_B_END    = 'тво_фам'
ANCHOR_ADDR  = 'відпустку буду проводити за адресою'


def paragraph_plain_text(p):
    return ''.join(r.text for r in p.runs)


def get_tvo_parts(donor_doc):
    paras = donor_doc.paragraphs
    a_idx = None
    for i, p in enumerate(paras):
        if TVO_A_MARK in paragraph_plain_text(p).lower():
            a_idx = i
            break
    if a_idx is None:
        return None, None
    b_start = None
    for i in range(a_idx + 1, len(paras)):
        if TVO_B_START in paragraph_plain_text(paras[i]).lower():
            b_start = i
            break
    if b_start is None:
        return None, None
    b_end = None
    for i in range(b_start, len(paras)):
        if TVO_B_END in paragraph_plain_text(paras[i]).lower().replace(' ', ''):
            b_end = i
            break
    if b_end is None:
        return None, None
    part_a = [copy.deepcopy(paras[a_idx]._p)]
    part_b = [copy.deepcopy(paras[i]._p) for i in range(b_start, b_end + 1)]
    return part_a, part_b


def find_addr_anchor(doc):
    for p in doc.paragraphs:
        if ANCHOR_ADDR in paragraph_plain_text(p).lower():
            return p
    return None


def is_applicant_signature(p):
    t = paragraph_plain_text(p).lower()
    return ('\u00abзван\u00bb' in t) and ('\u00abім\u00bb' in t or '\u00abiм\u00bb' in t) and ('\u00abфам\u00bb' in t)


def find_applicant_signature(doc):
    found = None
    for p in doc.paragraphs:
        if is_applicant_signature(p):
            found = p
    return found


def insert_tvo_block(target_doc, donor_doc):
    part_a, part_b = get_tvo_parts(donor_doc)
    if not part_a or not part_b:
        return False, 'donor'
    addr = find_addr_anchor(target_doc)
    if addr is None:
        return False, 'anchor_addr'
    sign = find_applicant_signature(target_doc)
    if sign is None:
        return False, 'anchor_sign'
    ref = sign._p
    for para_xml in part_b:
        ref.addnext(para_xml)
        ref = para_xml
    ref = addr._p
    for para_xml in part_a:
        ref.addnext(para_xml)
        ref = para_xml
    return True, None


AGE_CLAUSE_ANCHOR_NORM = 'спяніння'

AGE_CLAUSE_TEXT = (
    "Після закінчення особливого періоду, у разі досягнення військовослужбовцем "
    "граничного віку перебування на військовій службі, контракт припиняється "
    "(розривається), а військовослужбовець підлягає звільненню з військової служби за віком."
)


def insert_age_clause(doc):
    seen_ids = set()
    for p in collect_paragraphs(doc):
        pid = id(p._p)
        if pid in seen_ids:
            continue
        seen_ids.add(pid)
        text = normalize_for_anchor_(''.join(r.text for r in p.runs))
        if AGE_CLAUSE_ANCHOR_NORM in text:
            insert_paragraph_after(p, AGE_CLAUSE_TEXT)
            return True
    return False


LEADERSHIP_MATCH = {
    'komPolku': {
        'zvannia': 'майор', 'im': 'Максим', 'fam': 'ЗАЙЧЕНКО',
        'posPrefixes': ['Командир 1030', 'Командир військової частини А5273'],
    },
    'nachShtabu': {
        'zvannia': 'старший лейтенант', 'im': 'Євген', 'fam': 'СВИРИДОВ',
        'initialsFam': 'СВИРИДОВ',
        'posPrefixes': ['Начальник штабу'],
    },
    'komKorpusu': {
        'zvannia': 'бригадний генерал', 'im': 'Андрій', 'fam': 'БІЛЕЦЬКИЙ',
        'posPrefixes': ['Командир 3 армійського', 'Командир військової частини А5111'],
        'genitivePosPrefixes': ['Командира військової частини А5111'],
        'genitiveStandardText': 'бригадного генерала БІЛЕЦЬКОГО Андрія Євгенійовича',
    },
}

# ⚠️ «ТВО» = «тимчасово виконуючий обов'язки [КОГО?]» — граматично вимагає
# РОДОВОГО відмінка наступного слова, а не називного («ТВО командира»,
# а не «ТВО командир»). Це базове, безвиняткове правило для цих іменників
# чоловічого роду 2-ї відміни твердої групи (родовий = основа + «-а»).
# Джерело підтвердження класифікації слів: словникові правила бібліотеки
# shevchenko-ext-military (word-classifier-rules.json) — «командир»
# (закінчення «-ир») і «начальник» (закінчення «-ик») обидва класифікуються
# як masculine noun. Сам алгоритм відмінювання — у базовому пакеті
# shevchenko (JS/TS, окрема залежність), який Python-сервер не виконує;
# тому для цих конкретних, добре відомих слів відмінювання зроблено прямим
# точковим словником, а не викликом зовнішньої бібліотеки.
POSITION_HEAD_GENITIVE = {
    'командир': 'командира',
    'начальник': 'начальника',
}


def decline_position_phrase_to_genitive_(phrase):
    """Відмінює ЛИШЕ перше слово фрази (голову словосполучення) в родовий
       відмінок, залишаючи решту фрази («військової частини А5111»,
       «штабу», «1030 зенітного...») без змін — вона вже граматично
       узгоджена з головним словом і сама від відмінка голови не залежить."""
    parts = phrase.split(None, 1)
    if not parts:
        return phrase
    first = parts[0]
    rest = parts[1] if len(parts) > 1 else ''
    first_norm = normalize_for_anchor_(first)
    genitive = POSITION_HEAD_GENITIVE.get(first_norm)
    if genitive is None:
        return phrase  # невідоме слово — безпечний відкат, нічого не ламаємо
    if first[:1].isupper():
        genitive = genitive[0].upper() + genitive[1:]
    return genitive + (' ' + rest if rest else '')


def _ws(s):
    return r'[\s\u00A0]+'.join(re.escape(w) for w in s.split())


def regex_replace_with_highlight(paragraph, pattern, repl_func, highlight=True, touched_ids=None):
    if touched_ids is not None and id(paragraph._p) in touched_ids:
        return False

    runs = paragraph.runs
    if not runs:
        return False
    full = ''.join(r.text for r in runs)
    matches = list(pattern.finditer(full))
    if not matches:
        return False

    for m in reversed(matches):
        runs = paragraph.runs
        texts = [r.text for r in runs]
        offsets = []
        pos = 0
        for t in texts:
            offsets.append(pos)
            pos += len(t)

        s, e = m.start(), m.end()
        value = repl_func(m)

        first_i = None
        for i in range(len(texts)):
            if offsets[i] <= s < offsets[i] + len(texts[i]):
                first_i = i
                break
        last_i = None
        for j in range(len(texts)):
            if offsets[j] < e <= offsets[j] + len(texts[j]):
                last_i = j
        if first_i is None:
            continue
        if last_i is None:
            last_i = len(texts) - 1

        prefix = texts[first_i][:s - offsets[first_i]]
        suffix = texts[last_i][e - offsets[last_i]:]

        base_run = runs[first_i]
        base_run.text = prefix

        new_r_elem = copy.deepcopy(base_run._r)
        base_run._r.addnext(new_r_elem)
        new_run = Run(new_r_elem, paragraph)
        new_run.text = value
        if highlight:
            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        if first_i == last_i:
            suf_elem = copy.deepcopy(base_run._r)
            new_r_elem.addnext(suf_elem)
            suf_run = Run(suf_elem, paragraph)
            suf_run.text = suffix
        else:
            for k in range(first_i + 1, last_i):
                runs[k].text = ''
            runs[last_i].text = suffix

    if touched_ids is not None:
        touched_ids.add(id(paragraph._p))
    return True


def replace_pattern_everywhere(all_paragraphs, pattern, repl_func, highlight=True, touched_ids=None):
    changed = False
    for p in all_paragraphs:
        if regex_replace_with_highlight(p, pattern, repl_func, highlight=highlight, touched_ids=touched_ids):
            changed = True
    return changed


def replace_rank_in_same_row(all_tables, old_rank_norm, new_rank, name_pattern, touched_ids=None):
    changed = False
    seen_rows = set()
    for table in all_tables:
        for row in table.rows:
            row_id = id(row._tr)
            if row_id in seen_rows:
                continue
            seen_cells = set()
            cells = []
            for cell in row.cells:
                tcid = id(cell._tc)
                if tcid in seen_cells:
                    continue
                seen_cells.add(tcid)
                cells.append(cell)

            name_hit = any(name_pattern.search(cell_plain_text(c)) for c in cells)
            if not name_hit:
                continue
            seen_rows.add(row_id)

            for cell in cells:
                for p in cell.paragraphs:
                    if touched_ids is not None and id(p._p) in touched_ids:
                        continue
                    text = ''.join(r.text for r in p.runs)
                    if normalize_for_anchor_(text) == old_rank_norm:
                        runs = p.runs
                        if runs:
                            runs[0].text = new_rank
                            runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                            for r in runs[1:]:
                                r.text = ''
                            changed = True
                            if touched_ids is not None:
                                touched_ids.add(id(p._p))
    return changed


def _initials_of(im):
    im = (im or '').strip()
    return (im[0] + '.') if im else ''


def apply_leadership_substitution(doc, leadership_active):
    """Повертає діагностичний рядок (або None), якщо є що повідомити."""
    if not leadership_active:
        return None

    all_paragraphs = collect_all_paragraphs_incl_headers(doc)
    all_tables = walk_all_tables(doc)

    diag_lines = []
    full_text_norm = None
    touched_ids = set()

    for role_key, cfg in LEADERSHIP_MATCH.items():
        override = leadership_active.get(role_key)
        if not override:
            continue
        new_zv = (override.get('zvannia') or '').strip()
        new_im = (override.get('im') or '').strip()
        new_fam = (override.get('fam') or '').strip().upper()
        new_bat = (override.get('bat') or '').strip()
        new_genitive_text = (override.get('genitiveText') or '').strip()

        if not new_im or not new_fam:
            continue

        new_name = new_im + ' ' + new_fam
        new_initials = _initials_of(new_im)
        any_hit = False
        rank_combined_hit = False

        if cfg.get('initialsFam'):
            pat_init_full = re.compile(
                _ws(cfg['zvannia']) + r'[\s\u00A0]+' + _ws(cfg['initialsFam']) +
                r'[\s\u00A0]+[А-ЯІЇЄҐ]\.\s*[А-ЯІЇЄҐ]\.'
            )
            if replace_pattern_everywhere(all_paragraphs, pat_init_full,
                    lambda m: new_zv + ' ' + new_fam + ' ' + new_initials + '.',
                    touched_ids=touched_ids):
                any_hit = True
                rank_combined_hit = True

        pat_full = re.compile(
            _ws(cfg['zvannia']) + r'[\s\u00A0]+' + _ws(cfg['im']) + r'[\s\u00A0]+' + _ws(cfg['fam']),
            re.IGNORECASE
        )
        if replace_pattern_everywhere(all_paragraphs, pat_full, lambda m: new_zv + ' ' + new_name, touched_ids=touched_ids):
            any_hit = True
            rank_combined_hit = True

        pat_name = re.compile(_ws(cfg['im']) + r'[\s\u00A0]+' + _ws(cfg['fam']), re.IGNORECASE)
        name_replaced = replace_pattern_everywhere(all_paragraphs, pat_name, lambda m: new_name, touched_ids=touched_ids)
        if name_replaced:
            any_hit = True

        if cfg.get('initialsFam'):
            pat_init_bare = re.compile(_ws(cfg['initialsFam']) + r'[\s\u00A0]+[А-ЯІЇЄҐ]\.\s*[А-ЯІЇЄҐ]\.')
            if replace_pattern_everywhere(all_paragraphs, pat_init_bare,
                    lambda m: new_fam + ' ' + new_initials + '.', touched_ids=touched_ids):
                any_hit = True

        rank_row_hit = False
        if not rank_combined_hit and name_replaced:
            old_rank_norm = normalize_for_anchor_(cfg['zvannia'])
            new_name_pattern = re.compile(_ws(new_im) + r'[\s\u00A0]+' + _ws(new_fam))
            if replace_rank_in_same_row(all_tables, old_rank_norm, new_zv, new_name_pattern, touched_ids=touched_ids):
                rank_row_hit = True
                any_hit = True

        # ⚠️ Позиція — тепер «ТВО » + ГОЛОВНЕ СЛОВО фрази у РОДОВОМУ
        # відмінку (замість простого приліплення до називного).
        pos_hit = False
        for prefix in cfg.get('posPrefixes', []):
            pat_pos = re.compile(_ws(prefix))
            if replace_pattern_everywhere(
                all_paragraphs, pat_pos,
                lambda m: 'ТВО ' + decline_position_phrase_to_genitive_(m.group(0)),
                touched_ids=touched_ids
            ):
                pos_hit = True

        # Родовий контекст (напр. пункт контракту) — текст ТУТ уже в
        # родовому відмінку в самому шаблоні, тому просто додаємо «ТВО »,
        # без повторного відмінювання.
        genitive_hit = False
        if new_genitive_text and cfg.get('genitiveStandardText'):
            pat_gen_std = re.compile(_ws(cfg['genitiveStandardText']), re.IGNORECASE)
            if replace_pattern_everywhere(all_paragraphs, pat_gen_std,
                    lambda m: new_genitive_text, touched_ids=touched_ids):
                genitive_hit = True
                any_hit = True
            for prefix in cfg.get('genitivePosPrefixes', []):
                pat_gen_pos = re.compile(_ws(prefix))
                if replace_pattern_everywhere(all_paragraphs, pat_gen_pos,
                        lambda m: 'ТВО ' + m.group(0), touched_ids=touched_ids):
                    genitive_hit = True
                    any_hit = True

        if any_hit:
            rank_status = 'разом з іменем' if rank_combined_hit else ('в сусідній клітинці рядка' if rank_row_hit else 'НЕ ЗНАЙДЕНО окремо')
            diag_lines.append(
                'Роль «' + role_key + '» (' + new_zv + ' ' + new_name + '): ЗНАЙДЕНО і замінено (ПІБ: так, звання: ' +
                rank_status + ', посада: ' + ('так' if pos_hit else 'НІ') +
                (', родовий: ' + ('так' if genitive_hit else 'НІ') if new_genitive_text else '') + ').'
            )

    return '\n'.join(diag_lines) if diag_lines else None


def fill_document(docx_bytes, data_map, tvo_donor_bytes=None, add_tvo=False,
                  add_age_clause=False, warn_if_age_clause_missing=False,
                  leadership_active=None):
    doc = Document(io.BytesIO(docx_bytes))

    tvo_note = None
    if add_tvo:
        if tvo_donor_bytes is None:
            tvo_note = 'ТВО: не передано шаблон-донор (№2).'
        else:
            donor = Document(io.BytesIO(tvo_donor_bytes))
            ok, why = insert_tvo_block(doc, donor)
            if not ok:
                notes = {
                    'donor': 'ТВО: у доноровому шаблоні №2 не знайдено потрібні абзаци.',
                    'anchor_addr': 'ТВО: не знайдено абзац «Відпустку буду проводити за адресою…».',
                    'anchor_sign': 'ТВО: не знайдено підпис заявника («зван» «ім» «фам») для вставки блоку «Не заперечую».'
                }
                tvo_note = notes.get(why, 'ТВО: не вдалося вставити блок.')

    age_note = None
    if add_age_clause:
        found = insert_age_clause(doc)
        if not found and warn_if_age_clause_missing:
            age_note = ('Вік 45+: не знайдено абзац зі словом «сп\'яніння» '
                        'для вставки пункту 8 у цьому файлі.')

    apply_placeholders(doc, data_map)

    leadership_note = apply_leadership_substitution(doc, leadership_active)

    out = io.BytesIO()
    doc.save(out)

    notes = [n for n in (tvo_note, age_note, leadership_note) if n]
    note = '\n'.join(notes) if notes else None

    return out.getvalue(), note


@app.route('/', methods=['GET'])
def root():
    return 'DOCX filler service is running. VERSION: ' + APP_VERSION, 200


@app.route('/health', methods=['GET'])
def health():
    return 'ok | VERSION: ' + APP_VERSION, 200


@app.route('/generate', methods=['POST'])
def generate():
    try:
        job = request.get_json(force=True, silent=False)
    except Exception as e:
        return ('Bad JSON: ' + str(e), 400)
    if not job or 'people' not in job:
        return ('Missing "people"', 400)

    mem = io.BytesIO()
    with zipfile.ZipFile(mem, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('_ВЕРСІЯ_СЕРВІСУ.txt', ('Версія app.py, що обробила цей запит: ' + APP_VERSION).encode('utf-8'))

        for person in job.get('people', []):
            folder = person.get('folder', '') or ''
            data_map = person.get('data', {}) or {}
            add_tvo = bool(person.get('addTvo', False))
            add_age_clause = bool(person.get('addAgeClause', False))
            leadership_active = person.get('leadershipActive') or {}
            tvo_donor_b64 = person.get('tvoDonorB64')
            tvo_donor_bytes = base64.b64decode(tvo_donor_b64) if tvo_donor_b64 else None

            person_notes = []

            for f in person.get('files', []):
                out_name = f.get('outName', 'file.docx')
                b64 = f.get('templateB64', '')
                file_add_tvo = add_tvo and bool(f.get('allowTvo', True))
                looks_like_contract = 'контракт' in out_name.lower()
                if not b64:
                    continue
                try:
                    template_bytes = base64.b64decode(b64)
                    filled, note = fill_document(
                        template_bytes, data_map,
                        tvo_donor_bytes=tvo_donor_bytes,
                        add_tvo=file_add_tvo,
                        add_age_clause=add_age_clause,
                        warn_if_age_clause_missing=looks_like_contract,
                        leadership_active=leadership_active
                    )
                except Exception as e:
                    err = ('Помилка обробки: ' + str(e)).encode('utf-8')
                    path = (folder + '/' if folder else '') + out_name + '.ERROR.txt'
                    zf.writestr(path, err)
                    continue
                path = (folder + '/' if folder else '') + out_name
                zf.writestr(path, filled)
                if note:
                    person_notes.append(out_name + ':\n' + note)

            if person_notes:
                summary_name = (sanitize_for_filename_(folder) if folder else 'ХідЗаміни') + '_ХідЗаміни.txt'
                zf.writestr(summary_name, ('\n\n'.join(person_notes)).encode('utf-8'))

    return Response(
        mem.getvalue(),
        mimetype='application/zip',
        headers={'Content-Disposition': 'attachment; filename="result.zip"'}
    )


def sanitize_for_filename_(s):
    return re.sub(r'[\\/:*?"<>|]', '_', s or '').strip()


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
